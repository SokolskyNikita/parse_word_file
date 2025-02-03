import os
import sys
import time
import argparse
import logging
import json
from abc import ABC, abstractmethod
from pathlib import Path
from docx import Document
from docx.text.run import Run
from openai import OpenAI
from tqdm import tqdm
from tenacity import retry, stop_after_attempt, wait_exponential
from dataclasses import dataclass
from typing import Optional

DEFAULT_MODEL = "gpt-4o"
DEFAULT_MAX_TOKENS = 1000
DEFAULT_DELAY = 1.0
RETRY_COUNT = 3
RETRY_BACKOFF = 2.0
MIN_PARAGRAPH_LENGTH = 4

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    logging.error("Environment variable OPENAI_API_KEY not set.")
    sys.exit(1)

client = OpenAI(api_key=api_key)

@dataclass
class ProcessingConfig:
    model: str
    input_path: Path
    output_path: Path
    delay: float = DEFAULT_DELAY

class TextProcessor(ABC):
    def __init__(self, result_key: str, output_suffix: str) -> None:
        self.result_key = result_key
        self.output_suffix = output_suffix
    
    @abstractmethod
    def get_prompt(self, text: str) -> str:
        pass
    
    def extract_result(self, parsed_json: dict) -> Optional[str]:
        return parsed_json.get(self.result_key)
    
    def get_output_suffix(self) -> str:
        return self.output_suffix

class GrammarFixer(TextProcessor):
    def __init__(self) -> None:
        super().__init__(result_key="corrected_text", output_suffix="_grammar_fixed")
    
    def get_prompt(self, text: str) -> str:
        return (
            "Correct the following text into standard English. "
            f"Return your output as a JSON object with a key '{self.result_key}' containing the corrected text. "
            "Do not include any extra text or keys. Here is the text:\n\n"
            f"{text}"
        )

class Translator(TextProcessor):
    def __init__(self, target_language: str) -> None:
        self.target_language = target_language
        super().__init__(
            result_key="translated_text",
            output_suffix=f"_translated_to_{target_language.lower()}"
        )
    
    def get_prompt(self, text: str) -> str:
        return (
            f"Translate the following text into {self.target_language}. "
            f"Return your output as a JSON object with a key '{self.result_key}' containing the translated text. "
            "Maintain the same tone and formality as the original. "
            "Do not include any extra text or keys. Here is the text:\n\n"
            f"{text}"
        )

class DocumentProcessor:
    def __init__(self, config: ProcessingConfig, processor: TextProcessor) -> None:
        self.config = config
        self.processor = processor
    
    @retry(stop=stop_after_attempt(RETRY_COUNT), wait=wait_exponential(multiplier=RETRY_BACKOFF))
    def _call_openai_api(self, text: str) -> str:
        response = client.chat.completions.create(
            model=self.config.model,
            messages=[{"role": "user", "content": self.processor.get_prompt(text)}],  # type: ignore
            max_tokens=DEFAULT_MAX_TOKENS,
            frequency_penalty=0.0,
            response_format={"type": "json_object"}
        )
        
        if not (response.choices and response.choices[0].message and response.choices[0].message.content):
            return text
            
        try:
            parsed = json.loads(response.choices[0].message.content.strip())
            return self.processor.extract_result(parsed) or text
        except json.JSONDecodeError:
            logging.error("Failed to parse JSON response")
            return text
    
    def _apply_formatting(self, src_run: Run, tgt_run: Run) -> None:
        # Basic text properties
        if src_run.bold is not None:
            tgt_run.bold = src_run.bold
        if src_run.italic is not None:
            tgt_run.italic = src_run.italic
            
        # Font properties
        if src_run.font.double_strike is not None:
            tgt_run.font.double_strike = src_run.font.double_strike
        if src_run.font.all_caps is not None:
            tgt_run.font.all_caps = src_run.font.all_caps
        if src_run.font.small_caps is not None:
            tgt_run.font.small_caps = src_run.font.small_caps
        if src_run.font.name is not None:
            tgt_run.font.name = src_run.font.name
        if src_run.font.size is not None:
            tgt_run.font.size = src_run.font.size
        if src_run.font.highlight_color is not None:
            tgt_run.font.highlight_color = src_run.font.highlight_color
        if src_run.font.subscript is not None:
            tgt_run.font.subscript = src_run.font.subscript
        if src_run.font.superscript is not None:
            tgt_run.font.superscript = src_run.font.superscript
            
        # RGB color needs special handling due to nested property
        if src_run.font.color and src_run.font.color.rgb:
            tgt_run.font.color.rgb = src_run.font.color.rgb
    
    def process(self) -> None:
        doc_in = Document(str(self.config.input_path))
        doc_out = Document()
        
        paragraphs = [p for p in doc_in.paragraphs if p.text and len(p.text.strip()) >= MIN_PARAGRAPH_LENGTH]
        logging.info(f"Processing {len(paragraphs)} paragraphs...")
        
        for paragraph in tqdm(paragraphs, desc="Processing paragraphs"):
            new_para = doc_out.add_paragraph()
            new_run = new_para.add_run(self._call_openai_api(paragraph.text.strip()))
            
            for run in paragraph.runs:
                if run.text.strip():
                    self._apply_formatting(run, new_run)
                    break
            
            time.sleep(self.config.delay)
        
        doc_out.save(str(self.config.output_path))
        logging.info(f"Processed document saved to {self.config.output_path}")

def clean_language_name(lang: str) -> str:
    cleaned = ''.join(c for c in lang if c.isalpha()).lower().capitalize()
    if len(cleaned) < 3:
        raise ValueError(f"Language name '{lang}' is too short after cleaning. Must be at least 3 letters.")
    return cleaned

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Process a DOCX file using GPT models for grammar correction or translation."
    )
    parser.add_argument("input", type=str, help="Path to the input DOCX file")
    parser.add_argument("-o", "--output", type=str, help="Output DOCX file path")
    parser.add_argument("--model", type=str, default=DEFAULT_MODEL, help="OpenAI model to use")
    
    mode_group = parser.add_mutually_exclusive_group()
    mode_group.add_argument("-fg", "--fix-grammar", action="store_true", default=True,
                           help="Fix grammar (default)")
    mode_group.add_argument("-t", "--translate", type=str, metavar="LANG",
                           help="Translate to specified language")
    
    args = parser.parse_args()
    
    if args.translate:
        try:
            args.translate = clean_language_name(args.translate)
        except ValueError as e:
            parser.error(str(e))
    
    return args

def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    
    if not input_path.exists():
        logging.error(f"Input file '{input_path}' does not exist.")
        sys.exit(1)
    
    processor = Translator(args.translate) if args.translate else GrammarFixer()
    logging.info(f"Using processor: {processor.__class__.__name__}" + 
                (f" targeting language: {args.translate}" if args.translate else ""))
    
    output_path = Path(args.output) if args.output else input_path.with_name(
        f"{input_path.stem}{processor.get_output_suffix()}{input_path.suffix}"
    )
    
    config = ProcessingConfig(model=args.model, input_path=input_path, output_path=output_path)
    DocumentProcessor(config, processor).process()

if __name__ == "__main__":
    main()
