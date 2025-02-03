import os
import sys
import time
import argparse
import logging
import json
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openai import OpenAI  # Assuming you have a recent version that supports this interface
from tqdm import tqdm
from tenacity import retry, stop_after_attempt, wait_exponential

# Constants
FIX_GRAMMAR_TEMPLATE = (
    "Correct the following text into standard English. "
    "Return your output as a JSON object with a key 'corrected_text' containing the corrected text. "
    "Do not include any extra text or keys. Here is the text:\n\n%input%"
)
DEFAULT_TEMPLATE = FIX_GRAMMAR_TEMPLATE
# It is advisable to use a fully qualified model name (e.g. "gpt-4o-2024-08-06") if available.
DEFAULT_MODEL = "gpt-4o"
DEFAULT_MAX_TOKENS = 1000
DEFAULT_DELAY = 1.0
RETRY_COUNT = 3
RETRY_BACKOFF = 2.0

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s"
)

# Updated OpenAI API Setup – if using the new client interface:
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Argument Parsing
def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Process a DOCX file using GPT-4o for grammar correction."
    )
    parser.add_argument("input", type=str, help="Path to the input DOCX file.")
    parser.add_argument(
        "-o",
        "--output",
        type=str,
        help="Output DOCX file path (default: <input>_parsed.docx).",
    )
    parser.add_argument(
        "--model",
        type=str,
        default=DEFAULT_MODEL,
        help=f"OpenAI model to use (default: {DEFAULT_MODEL}).",
    )
    return parser.parse_args()

# Prompt Formatting
def format_prompt(text: str, template: str = DEFAULT_TEMPLATE) -> str:
    return template.replace("%input%", text)

# OpenAI API Interaction with JSON parsing using tenacity for retries
@retry(
    stop=stop_after_attempt(RETRY_COUNT),
    wait=wait_exponential(multiplier=RETRY_BACKOFF),
)
def call_openai_api(prompt: str, model: str) -> str:
    # Create properly formatted message object
    messages = [
        {
            "role": "user",
            "content": prompt,
        }
    ]

    # Enforce JSON response mode by adding response_format parameter
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        max_tokens=DEFAULT_MAX_TOKENS,
        frequency_penalty=0.0,
        response_format={"type": "json_object"}
    )

    # Handle potential None response
    if not response.choices or not response.choices[0].message or not response.choices[0].message.content:
        logging.error("Empty or invalid response from API")
        return prompt

    output_text = response.choices[0].message.content.strip()
    
    # Try to parse JSON output
    try:
        parsed = json.loads(output_text)
        if "corrected_text" in parsed:
            return parsed["corrected_text"]
        logging.error("JSON missing 'corrected_text' key; using full output.")
        return output_text
    except json.JSONDecodeError:
        logging.error("Failed to parse JSON; returning raw output.")
        return output_text

def process_paragraph_text(text: str, template: str, model: str) -> str:
    prompt = format_prompt(text, template)
    try:
        return call_openai_api(prompt, model)
    except Exception as e:
        logging.error(f"API call failed after retries: {e}")
        return text

# Formatting Utilities: copy all relevant font properties
def apply_formatting(src_run: Run, tgt_run: Run) -> None:
    tgt_run.bold = src_run.bold
    tgt_run.italic = src_run.italic
    tgt_run.font.double_strike = src_run.font.double_strike
    tgt_run.font.all_caps = src_run.font.all_caps
    tgt_run.font.small_caps = src_run.font.small_caps
    tgt_run.font.name = src_run.font.name
    tgt_run.font.size = src_run.font.size
    if src_run.font.color and src_run.font.color.rgb:
        tgt_run.font.color.rgb = src_run.font.color.rgb
    tgt_run.font.highlight_color = src_run.font.highlight_color
    tgt_run.font.subscript = src_run.font.subscript
    tgt_run.font.superscript = src_run.font.superscript

def process_paragraph(paragraph: Paragraph, template: str, model: str) -> str:
    txt = paragraph.text.strip()
    return (
        txt
        if not txt or len(txt) <= 3
        else process_paragraph_text(txt, template, model)
    )

def process_document(
    input_path: Path, output_path: Path, template: str, model: str
) -> None:
    doc_in = Document(str(input_path))
    doc_out = Document()
    paras = [p for p in doc_in.paragraphs if p.text and len(p.text.strip()) > 3]
    logging.info(f"Processing {len(paras)} paragraphs...")
    for p in tqdm(paras, desc="Processing paragraphs"):
        corr = process_paragraph(p, template, model)
        new_para = doc_out.add_paragraph()
        new_run = new_para.add_run(corr)
        if p.runs:
            apply_formatting(p.runs[0], new_run)
        time.sleep(DEFAULT_DELAY)
    doc_out.save(str(output_path))
    logging.info(f"Saved processed document to {output_path}")

def main() -> None:
    args = parse_args()
    inp = Path(args.input)
    if not inp.exists():
        logging.error(f"Input file {inp} does not exist.")
        sys.exit(1)
    out = (
        Path(args.output)
        if args.output
        else inp.with_name(f"{inp.stem}_parsed{inp.suffix}")
    )
    process_document(inp, out, FIX_GRAMMAR_TEMPLATE, args.model)

if __name__ == "__main__":
    main()
